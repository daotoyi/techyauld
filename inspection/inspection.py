'''
Description: 
Author: daoyi
Date: 2021-12-03 21:00:09
LastEditors: daoyi
LastEditTime: 2021-12-08 02:21:17
'''
#-*- coding : utf-8 -*-
# coding: unicode_escape

from typing import DefaultDict
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Cm, Pt, Length, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from add_float_picture import add_float_picture
# from docx.enum.dml import MSO_THEME_COLOR

import xlwings as xw

from win32com.client import Dispatch

from copy import deepcopy
from pathlib import Path
import datetime
import os, sys
import json
import random
from interval import Interval

from gooey import Gooey, GooeyParser
import argparse

import pysnooper as snp
import logging

# %(filename)s [line:%(lineno)d] \
logging.basicConfig(level=logging.INFO, format='[%(asctime)s] \
[line:%(lineno)d] [%(levelname)s] : %(message)s',  datefmt='%Y-%m-%d(%a) %H:%M:%S')

CURRENT_PATH= os.getcwd()   # str
logging.info(f'working directory ==>> {CURRENT_PATH}')

USERSPACE_PATH = CURRENT_PATH + "/userspace/"
DATE = datetime.datetime.now().strftime('%Y-%m-%d')


def enum_file(filePath, type_) -> list:
    '''
    para
    :type_: .doc or .pdf
    '''
    list_file = []
    list_tmp = os.listdir(filePath)

    for i in range(0,len(list_tmp)):
        tmp = os.path.join(filePath, list_tmp[i])
        if os.path.isdir(tmp):
            list_file.extend(enum_file(tmp, type_))
        if os.path.isfile(tmp):
            list_file.append(tmp)

    object_file = []
    for i in list_file:
        (file_name, extension) = os.path.splitext(i)
        if extension == type_:
            (path, file) = os.path.split(i)
            object_file.append(file)

    return object_file

# @snp.snoop(depth=1, prefix="pdf: ")
def docx_to_pdf(source_file: str, output_file: str):
    word = Dispatch("Word.Application")
    word.Visible = 0  # run in backend
    word.DisplayAlerts = 0
    doc = word.Documents.Open(source_file)
    doc.SaveAs(output_file + ".pdf", 17)  # txt=4, html=10, docx=16， pdf=17
    doc.Close()
    word.Quit()

# @snp.snoop(depth=1, prefix="handle: ")
def sort_table(xlsfile, customer, start_date, end_date) -> list:
    workbook = xw.Book(xlsfile)

    sheets_list = workbook.sheets
    # select_custom = '许继软件'
    select_custom = customer
    select_attribute = '销售'
    new_sheet = workbook.sheets.add(select_custom)
    select_interval = Interval(
        datetime.datetime.strptime(start_date, '%Y-%m-%d'), 
        datetime.datetime.strptime(end_date, '%Y-%m-%d')
        )

    range_value_list = []
    def readrange(sheet, nrows):
        for i in range(2, nrows):
            #单个表格字符串
            # sheet_sale = "L"+str(i)
            sheet_date = "G"+str(i)
            sheet_custome = "I"+str(i)
            sheet_attribure = "J"+str(i)

            #整行表格字符串
            select_sheet = "F"+str(i)+":"+"L"+str(i)

            # select_sheet_sale = sheet.range(sheet_sale).value
            select_sheet_date = sheet.range(sheet_date).value
            select_sheet_custome = sheet.range(sheet_custome).value
            select_sheet_attribute = sheet.range(sheet_attribure).value

            logging.debug(type(select_sheet_date))
            logging.debug(type(select_interval))

            if select_sheet_date in select_interval and \
               select_sheet_custome == select_custom and \
               select_sheet_attribute == select_attribute:
                str_value_row = sheet.range(select_sheet).value
                range_value_list.append(str_value_row) 

    def readrange_default(sheet, nrows):
        for i in range(2, nrows):
            sheet_custome = "I"+str(i)
            sheet_attribure = "J"+str(i)

            select_sheet = "F"+str(i)+":"+"L"+str(i)

            select_sheet_custome = sheet.range(sheet_custome).value
            select_sheet_attribute = sheet.range(sheet_attribure).value

            if select_sheet_custome == select_custom and \
                select_sheet_attribute == select_attribute:
                str_value_row = sheet.range(select_sheet).value
                range_value_list.append(str_value_row) 


    for sheet in sheets_list:
        rng = sheet.range('a1').expand('table')
        nrows = rng.rows.count
        logging.info('Organize Sheet Data ==>> Waiting...')
        readrange_default(sheet, nrows) if start_date == DATE and end_date == DATE \
            else readrange(sheet, nrows)

    new_sheet.range("A1:G1").value = ["主机序列号","出库日期","出库单号","客户名称","出货属性","订单号","经手人"]
    row_num = 1
    for row in range_value_list:
        row_num += 1
        new_sheet_row = "A"+str(row_num)+":"+"G"+str(row_num)
        new_sheet.range(new_sheet_row).value = row
        logging.info(f'Write Sheet<{select_custom}> ==>> ' + f'{len(range_value_list)}' + '/' + f'{row_num-1}')

    ## -------------------------------------------
    sht = workbook.sheets(select_custom)
    info = sht.used_range
    nrows = info.last_cell.row
    ncolumns = info.last_cell.column
    logging.info(f'total rows: {nrows}\n')
    if nrows == 1:
        logging.info("error, no select data, please set again.")
        sys.exit()

    psnlist = sht.range(f'a1:a{nrows}').value
    datelist= sht.range(f'b1:b{nrows}').value

    return list(zip(psnlist, datelist))

def handle(data_list, save_folder, file_name_formate):
    try:
        with open(USERSPACE_PATH + 'match.json') as j:
            match = json.load(j)
    except:
        match = {
            "T4689":"WYD-811",
            "T4582":"WYD-811",
            "T4471":"MCE-812",
            "T4496":"MCE-812",
            "T4794":"SPMU-852",
            "T4254":"WEA-852",
            "TR529":"WEA-852",
            "T5DB6":"PAC-8581"
        }

    num = 1
    # custome_type = match[psn[:5]] if psn[:5] in match.keys() else continue
    data = data_list[1:]
    for psn, date in data:
        logging.debug(date)
        date = date.strftime('%Y-%m-%d')
        if psn[:5] in match.keys():
            custome_type = match[psn[:5]]
        else:
            continue

        template = USERSPACE_PATH + "template/" + custome_type + ".docx"
        if not os.path.exists(template):
            template = USERSPACE_PATH + "template/" + "template.docx"

        wordfile = Document(template)
        wordfile_copy = deepcopy(wordfile)  # 防止原文件被篡改，deepcopy 

        for paragraphs in wordfile_copy.paragraphs:
            for run in paragraphs.runs:
                # run.text = run.text.replace("PSN", str(psn))
                # run.text = run.text.replace("TYPE", str(custome_type))
                run.text = run.text.replace("DATE", str(date))

        for table in wordfile_copy.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace("PSN", str(psn))
                    cell.text = cell.text.replace("TYPE", str(custome_type))
                    cell.text = cell.text.replace("DATE", str(date))
                    cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell.paragraphs[0].runs[0].font.size = Pt(13)

        # wordfile_copy.add_picture(CURRENT_PATH + '/template/signature.png',width=Inches(2))
        par = wordfile_copy.add_paragraph()
        posix = random.randint(5, 13)    # better 12
        posiy = random.randint(16, 20)   # better 18
        add_float_picture(
            par, 
            USERSPACE_PATH + 'template/signature.png', 
            width=Inches(2.5), 
            pos_x=Cm(posix), 
            pos_y=Cm(posiy)
        )

        wordfile_copy.save(f'{save_folder}/{psn}.docx') if file_name_formate == "PSN" \
            else wordfile_copy.save(f'{save_folder}/{custome_type}@{psn}.docx')
        logging.info('outputdocx ==>> ' + f'{len(data)}' + '/' + f'{num}')
        num = num +1
    logging.info('outputdocx =======>> ' + f'{len(data) - num + 1}' + " files not contained.") 

# @snp.snoop(depth=1, prefix="inspection: ")
def custom_inspection(file, customer, file_name_formate, start_date, end_date):

    # xlsfile = enum_file(USERSPACE_PATH, '.xlsx')[0] if not file else file
    try:
        xlsfile = USERSPACE_PATH + enum_file(USERSPACE_PATH, '.xlsx')[0] if not os.path.exists(file) else file
        logging.debug(xlsfile)
    except:
        logging.info("===>> userspace directory not exist psn.xlsx <<===")
        return

    save_docx_folder = Path(USERSPACE_PATH + f'{DATE}@outputdocx')
    save_docx_folder.mkdir(parents=True, exist_ok=True)
    handle(sort_table(xlsfile, customer, start_date, end_date), save_docx_folder, file_name_formate)
    docx_list = enum_file(save_docx_folder, ".docx")
    logging.debug(docx_list)

    save_pdf_folder = Path(USERSPACE_PATH + f'{DATE}@outputpdf')
    save_pdf_folder.mkdir(parents=True, exist_ok=True)
    num = 1
    for docx in docx_list:
        source_file = os.path.join(save_docx_folder,docx)
        output_file = os.path.join(save_pdf_folder, docx[:-5])
        docx_to_pdf(source_file, output_file)

        logging.info('docx2pdf ==>> ' + f'{len(docx_list)}' + '/' + f'{num}')
        num = num +1

    # from docx2pdf import convert
    # convert("C:/Users/Test/Desktop/out.docx", "C:/Users/Test/Desktop/output.pdf")
    logging.info('===>>> done <<<===')

@Gooey(
    richtext_controls=True,                 # 打开终端对颜色支持
    program_name="许继质检单",
    encoding='utf-8',                       # 设置编码格式utf-8，打包的时候遇到问题
    # optional_cols=2,
    default_size=(650,500),
    language='chinese',
    progress_regex=r"^progress: (\d+)%$"    # 正则，用于模式化运行时进度信息
)
def main():
    parser = GooeyParser(description="根据模板自动生成客户质检单\n<whshi@techyauld.com>")
    parser.add_argument('Customer', help="选择客户", default="许继软件")
    parser.add_argument('FileName', help="保存文件名", widget="Dropdown", choices=['PSN','TYPE@PSN'], default="TYPE@PSN") 
    parser.add_argument('StartDate', help="开始日期", widget="DateChooser", default=f'{DATE}')
    parser.add_argument('EndDate', help="截止日期", widget="DateChooser", default=f'{DATE}')
    parser.add_argument('Path', help="文件路径", widget="FileChooser", default="userspace/psn.xlsx (default)") 

    args = parser.parse_args()          # 接收界面传递的参数
    print(args, flush=True)             # flush=True在打包的时候会用到

    # --------------------------------------------------------------------
    parser_arg = argparse.ArgumentParser(description="My Cool Gooey App!")
    parser_arg.add_argument('Customer', help="客户", default="许继软件")
    # args_arg = parser_arg.parse_args()    # 若同时parser_arg解析，则程序报错 
    # --------------------------------------------------------------------

    logging.debug(f'parse args: {args.FileName}')
    logging.debug(type(args.StartDate))
    custom_inspection(
        args.Path, 
        args.Customer, 
        args.FileName, 
        args.StartDate,
        args.EndDate
        )


if __name__ == "__main__":
    main()
    # custom_inspection("")
